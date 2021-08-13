# -*- coding: utf-8 -*-
import os
import docx
import sys

# file path
path = u'D:\\Documents\\chromeDownload\\CIMS_SOPs\\Dept. 70 - System Dev'
# auto create log
tlog =  path + u'\\replaceFilesList.txt'
err_log = path + u'\\replaceErrorList.txt'

if sys.getdefaultencoding() != 'utf-8':
    reload(sys)
    sys.setdefaultencoding('utf-8')
# create two log
def log(text):
    with open( err_log,"a+" ) as f:
        f.write(text)
        f.write('\n')
def log2(text):
    with open( tlog,"a+" ) as f:
        f.write(text)
        f.write('\n')


# Replace content (document name, old content, new content)
def info_update(doc,old_info,new_info):
    # Replace all text in the document
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace(old_info,new_info)
    # Replace the contents of the table in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info,new_info)


                
def thr(old_info,new_info):
    # Traverse the docx documents in the directory
    for parent, dirnames, filenames in os.walk(path):
        for fn in filenames:
            filedir = os.path.join(parent, fn)
            if fn.endswith('.docx'):
                try:
                    # Define document path
                    doc = docx.Document(filedir)
                    # Change the font to Times New Roman
                    doc.styles['Normal'].font.name = u'Times New Roman'
                    # Call function to modify the content of the document
                    info_update(doc,old_info,new_info)
                    # Sace document
                    doc.save(filedir)
                    # Write to the modification log
                    log2(filedir + ' #### Modification is complete')
                    print(filedir + ' #### Modification is complete')
                except Exception as e:
                    # Writeto the  modification failure log
                    log(filedir)


if __name__ == '__main__':
    thr('Brightech International','CIMS Global')
    thr('brightech','CIMS Global')
    thr('Brightech','CIMS Global')
   
    print('----Replace all documents----')
